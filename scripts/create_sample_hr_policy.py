#!/usr/bin/env python3
"""Create a sample HR policy docx for extraction validation."""
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Run: pip install python-docx")
    raise

def main():
    doc = Document()
    doc.add_heading("International Relocation Policy", 0)
    doc.add_paragraph("Version 1.0 | Effective 2026-01-01")

    doc.add_heading("6.1 Temporary Housing", level=2)
    doc.add_paragraph(
        "Temporary Housing is provided for up to 60 days for B1/B2 bands, 90 days for B3, and 120 days for B4. "
        "Monthly cap: Oslo NOK 30,000, New York USD 6,000, London GBP 4,500, Singapore SGD 7,000."
    )

    doc.add_heading("6.2 Shipment of Household Goods", level=2)
    doc.add_paragraph(
        "Shipment of household goods: one full container for Permanent and Long-Term assignments. "
        "Short-Term: 500 kg. Covers packing, transport, and customs clearance."
    )

    doc.add_heading("6.3 Education Support", level=2)
    doc.add_paragraph(
        "Education support: up to 75% of tuition for international school placement. "
        "Applies to Permanent and Long-Term assignments. Bands B1-B4."
    )

    doc.add_heading("6.4 Visa & Immigration", level=2)
    doc.add_paragraph(
        "Visa and work permit support: employer-sponsored. Immigration support includes residence permit, "
        "dependent visas, and relocation agency support."
    )

    doc.add_heading("6.5 Travel to Host Location", level=2)
    doc.add_paragraph(
        "Travel to host location: economy class flights for employee and dependents. "
        "One-way for Permanent, round-trip for Long-Term and Short-Term."
    )

    doc.add_heading("6.6 Scouting Trip", level=2)
    doc.add_paragraph(
        "Pre-assignment visit (scouting trip): one trip, up to 5 days. Cap: USD 3,000. "
        "Covers flights, accommodation, and local transport."
    )

    doc.add_heading("6.7 Settling-in Allowance", level=2)
    doc.add_paragraph(
        "Settling-in allowance: lump sum based on family size. Permanent: 3 months gross. "
        "Long-Term: 2 months. Short-Term: 1 month."
    )

    doc.add_heading("6.8 Tax Assistance", level=2)
    doc.add_paragraph(
        "Tax assistance and tax equalization available for Permanent and Long-Term assignments. "
        "Covers preparation, filing, and advisory."
    )

    doc.add_heading("6.9 Spousal Support", level=2)
    doc.add_paragraph(
        "Spousal support: career transition assistance up to USD 5,000. "
        "Partner support includes job search and networking."
    )

    doc.add_heading("6.10 Language & Cultural Training", level=2)
    doc.add_paragraph(
        "Language and cultural training: up to 40 hours per family. "
        "Integration support covers local orientation and cultural briefing."
    )

    doc.add_heading("6.11 Repatriation", level=2)
    doc.add_paragraph(
        "Repatriation: return shipment and return travel at end of assignment. "
        "Same entitlements as outbound relocation."
    )

    out_dir = Path(__file__).resolve().parent.parent / "backend" / "tests" / "fixtures"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "HR_policy_2026.docx"
    doc.save(out_path)
    print(f"Created: {out_path}")

if __name__ == "__main__":
    main()
