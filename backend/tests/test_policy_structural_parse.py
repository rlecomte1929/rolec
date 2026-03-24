"""Structural parse facade: native path and optional-backend placeholders."""
from __future__ import annotations

import io
import os
import sys
import unittest

_REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
if _REPO_ROOT not in sys.path:
    sys.path.insert(0, _REPO_ROOT)

from backend.services.policy_document_clauses import extract_lines_with_pages
from backend.services.policy_structural_parse import (
    StructuralParseBackend,
    parse_policy_document_to_elements,
)


def _minimal_docx_bytes() -> bytes:
    try:
        from docx import Document  # type: ignore
    except ImportError as e:
        raise unittest.SkipTest("python-docx not installed") from e

    doc = Document()
    doc.add_paragraph("Long term assignment policy summary")
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Benefit"
    table.rows[0].cells[1].text = "Detail"
    table.rows[0].cells[2].text = "2.1"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class PolicyStructuralParseTests(unittest.TestCase):
    def test_native_matches_extract_lines_with_pages(self) -> None:
        data = _minimal_docx_bytes()
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        direct, err_d = extract_lines_with_pages(data, mime)
        self.assertIsNone(err_d)
        fac, err_f = parse_policy_document_to_elements(
            data, mime, backend=StructuralParseBackend.NATIVE, fallback_on_error=True
        )
        self.assertIsNone(err_f)
        self.assertEqual(len(direct), len(fac))
        for a, b in zip(direct, fac):
            self.assertEqual(a.get("text"), b.get("text"))
            self.assertEqual(a.get("page"), b.get("page"))
            self.assertEqual(a.get("is_table_row"), b.get("is_table_row"))
            self.assertEqual(b.get("structural_source"), "native")

    def test_docling_placeholder_falls_back_when_enabled(self) -> None:
        data = _minimal_docx_bytes()
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        fac, err = parse_policy_document_to_elements(
            data,
            mime,
            backend=StructuralParseBackend.DOCLING,
            fallback_on_error=True,
        )
        self.assertIsNone(err)
        self.assertGreater(len(fac), 0)
        self.assertEqual(fac[0].get("structural_source"), "native")

    def test_docling_no_fallback_returns_error(self) -> None:
        fac, err = parse_policy_document_to_elements(
            b"%PDF-1.4",
            "application/pdf",
            backend=StructuralParseBackend.DOCLING,
            fallback_on_error=False,
        )
        self.assertIsNotNone(err)
        self.assertEqual(fac, [])


if __name__ == "__main__":
    unittest.main()
