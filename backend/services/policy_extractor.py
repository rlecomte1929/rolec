import io
import re
from typing import Any, Dict, List, Tuple

from datetime import datetime


CATEGORY_KEYWORDS: Dict[str, List[str]] = {
    "housing": ["temporary housing", "housing", "rental", "deposit", "accommodation"],
    "movers": ["shipment", "household goods", "movers", "moving", "freight", "shipping"],
    "schools": ["education", "school", "tuition", "childcare"],
    "immigration": ["visa", "immigration", "work permit", "residence permit"],
    "travel": ["travel", "flight", "airfare", "scouting trip", "pre-assignment visit"],
    "settling_in": ["settling-in", "settling in", "allowance"],
    "tax": ["tax assistance", "tax equalization", "tax"],
    "spouse": ["spousal support", "spouse", "partner support"],
    "integration": ["language", "cultural", "integration", "training"],
    "repatriation": ["repatriation", "return shipment", "return travel"],
    "home_sale": ["home sale", "home purchase", "property sale", "property purchase"],
}


BENEFIT_KEYS: List[Tuple[str, str, List[str], str]] = [
    ("temporary_housing", "Temporary housing duration", ["temporary housing", "temporary accommodation"], "housing"),
    ("rental_deposit", "Rental deposit support", ["deposit", "rental deposit"], "housing"),
    ("shipment", "Shipment of household goods", ["shipment", "household goods", "moving"], "movers"),
    ("education_support", "Education support", ["education", "school", "tuition"], "schools"),
    ("visa_support", "Visa & immigration support", ["visa", "immigration", "work permit"], "immigration"),
    ("travel_host", "Travel to host location", ["travel", "flight", "airfare"], "travel"),
    ("settling_in_allowance", "Settling-in allowance", ["settling-in", "settling in allowance"], "settling_in"),
    ("tax_assistance", "Tax assistance", ["tax assistance", "tax equalization"], "tax"),
    ("spousal_support", "Spousal support", ["spousal", "partner support"], "spouse"),
    ("language_training", "Language/cultural training", ["language", "cultural", "training"], "integration"),
    ("repatriation", "Repatriation", ["repatriation", "return shipment"], "repatriation"),
    ("scouting_trip", "Scouting trip", ["scouting trip", "pre-assignment visit"], "travel"),
]


def _normalize_text(lines: List[str]) -> List[str]:
    cleaned = []
    for line in lines:
        if not line:
            continue
        s = re.sub(r"\s+", " ", line.strip())
        if s:
            cleaned.append(s)
    return cleaned


def _extract_text_from_docx(data: bytes) -> List[str]:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise RuntimeError("python-docx is required for .docx extraction") from exc
    doc = Document(io.BytesIO(data))
    lines: List[str] = []
    for p in doc.paragraphs:
        lines.append(p.text or "")
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text]
            if cells:
                lines.append(" | ".join(cells))
    return _normalize_text(lines)


def _extract_text_from_pdf(data: bytes) -> List[str]:
    try:
        import pdfplumber  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise RuntimeError("pdfplumber is required for .pdf extraction") from exc
    lines: List[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            lines.extend(text.splitlines())
            for table in page.extract_tables() or []:
                for row in table:
                    cells = [c for c in row if c]
                    if cells:
                        lines.append(" | ".join(cells))
    return _normalize_text(lines)


def _guess_meta(lines: List[str]) -> Dict[str, Any]:
    title = ""
    version = ""
    effective_date = None
    for line in lines[:40]:
        if not title and "policy" in line.lower():
            title = line
        if not version:
            m = re.search(r"(version|v\.)\s*([0-9]+(?:\.[0-9]+)?)", line, re.I)
            if m:
                version = m.group(2)
        if not effective_date:
            m = re.search(r"(effective|valid from)\s*[:\-]?\s*([0-9]{4}-[0-9]{2}-[0-9]{2})", line, re.I)
            if m:
                effective_date = m.group(2)
    return {
        "title": title or "Relocation Policy",
        "version": version or None,
        "effective_date": effective_date,
    }


def _extract_limits(text: str) -> Dict[str, Any]:
    limits: Dict[str, Any] = {}
    day_match = re.search(r"(\d{1,3})\s*(days|day)\b", text, re.I)
    if day_match:
        limits["days"] = int(day_match.group(1))
    percent_match = re.search(r"(\d{1,3})\s*%", text)
    if percent_match:
        limits["percent"] = int(percent_match.group(1))
    amounts: Dict[str, float] = {}
    for m in re.finditer(r"([A-Z]{2,3})\s?([0-9][0-9,\.]+)", text):
        cur = m.group(1).upper()
        val = float(m.group(2).replace(",", ""))
        amounts[cur] = val
    for m in re.finditer(r"([0-9][0-9,\.]+)\s?(USD|EUR|NOK|GBP|SGD)", text, re.I):
        cur = m.group(2).upper()
        val = float(m.group(1).replace(",", ""))
        amounts[cur] = val
    if amounts:
        key = "monthly_cap" if re.search(r"month", text, re.I) else "cap"
        limits[key] = amounts
    return limits


def _extract_eligibility(text: str) -> Dict[str, Any]:
    bands = sorted(set(re.findall(r"\bB[1-4]\b", text)))
    assignment_types = []
    for k in ["Permanent", "Long-Term", "Short-Term"]:
        if re.search(k, text, re.I):
            assignment_types.append(k.lower().replace("-", "_"))
    elig: Dict[str, Any] = {}
    if bands:
        elig["bands"] = bands
    if assignment_types:
        elig["assignment_types"] = assignment_types
    return elig


def extract_policy_from_bytes(file_bytes: bytes, file_type: str) -> Dict[str, Any]:
    if file_type == "docx":
        lines = _extract_text_from_docx(file_bytes)
    elif file_type == "pdf":
        lines = _extract_text_from_pdf(file_bytes)
    else:
        raise ValueError("Unsupported file type")

    meta = _guess_meta(lines)
    benefits: List[Dict[str, Any]] = []
    seen_keys: set = set()

    for line in lines:
        lower = line.lower()
        for key, label, keywords, category in BENEFIT_KEYS:
            if key in seen_keys:
                continue
            if any(k in lower for k in keywords):
                elig = _extract_eligibility(line)
                limits = _extract_limits(line)
                benefits.append(
                    {
                        "service_category": category,
                        "benefit_key": key,
                        "benefit_label": label,
                        "eligibility": elig or None,
                        "limits": limits or None,
                        "notes": None,
                        "source_section": None,
                        "source_quote": line[:200],
                        "confidence": 0.6 if limits or elig else 0.4,
                    }
                )
                seen_keys.add(key)

    # Add fallback entries for critical categories if found in document body
    if not benefits:
        for line in lines[:80]:
            for category, keywords in CATEGORY_KEYWORDS.items():
                if any(k in line.lower() for k in keywords):
                    benefit_key = f"{category}_support"
                    if benefit_key in seen_keys:
                        continue
                    benefits.append(
                        {
                            "service_category": category,
                            "benefit_key": benefit_key,
                            "benefit_label": f"{category.replace('_', ' ').title()} support",
                            "eligibility": None,
                            "limits": _extract_limits(line) or None,
                            "notes": None,
                            "source_section": None,
                            "source_quote": line[:200],
                            "confidence": 0.3,
                        }
                    )
                    seen_keys.add(benefit_key)

    return {
        "policy_meta": meta,
        "benefits": benefits,
        "extracted_at": datetime.utcnow().isoformat(),
    }
