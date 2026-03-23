"""
Policy document clause segmentation: split raw text into structured clause blocks
with source traceability. No normalization to final policy rules here.

`normalized_hint_json` is **Layer 1** (heuristic hints). It feeds the normalization
transform into **Layer 2** (`policy_benefit_rules`, etc.) but is not itself a machine
decision — see docs/policy/metadata-vs-decision-layer.md.
"""
from __future__ import annotations

import logging
import re
from typing import Any, Dict, List, Optional, Tuple

log = logging.getLogger(__name__)

# Clause types
CLAUSE_SCOPE = "scope"
CLAUSE_ELIGIBILITY = "eligibility"
CLAUSE_BENEFIT = "benefit"
CLAUSE_EXCLUSION = "exclusion"
CLAUSE_APPROVAL_RULE = "approval_rule"
CLAUSE_EVIDENCE_RULE = "evidence_rule"
CLAUSE_TAX_RULE = "tax_rule"
CLAUSE_DEFINITION = "definition"
CLAUSE_LIFECYCLE_RULE = "lifecycle_rule"
CLAUSE_UNKNOWN = "unknown"

# Section headings that map to clause sections (case-insensitive)
SECTION_HEADINGS = [
    "scope and exclusions",
    "scope",
    "exclusions",
    "family status",
    "accompanying family",
    "social coverage",
    "social security",
    "taxation",
    "tax equalization",
    "moving services",
    "household goods",
    "accommodation",
    "housing",
    "schooling",
    "education",
    "tuition",
    "home leave",
    "support to partner",
    "spousal support",
    "tax equalization",
    "responsibility matrix",
    "long term assignment",
    "short term assignment",
    "policy summary",
    "eligibility",
    "definitions",
]

# Clause-type keyword mapping (phrase -> clause_type)
CLAUSE_TYPE_SIGNALS: List[Tuple[List[str], str]] = [
    (["scope", "applies to", "in scope", "out of scope", "coverage applies"], CLAUSE_SCOPE),
    (["eligibility", "eligible", "qualify", "qualification"], CLAUSE_ELIGIBILITY),
    (["exclusion", "excluded", "not covered", "ineligible", "does not apply"], CLAUSE_EXCLUSION),
    (["approval", "pre-approval", "prior approval", "requires approval", "hr approval"], CLAUSE_APPROVAL_RULE),
    (["evidence", "receipt", "invoice", "documentation required", "submit receipts"], CLAUSE_EVIDENCE_RULE),
    (["tax equalization", "hypothetical tax", "tax protection", "tax assistance"], CLAUSE_TAX_RULE),
    (["definition", "defined as", "means ", "refers to"], CLAUSE_DEFINITION),
    (["repatriation", "return shipment", "end of assignment", "lifecycle"], CLAUSE_LIFECYCLE_RULE),
    (["housing", "allowance", "moving", "school", "tuition", "home leave", "mobility premium"], CLAUSE_BENEFIT),
]

# Benefit key candidates (normalized keys for downstream)
BENEFIT_KEY_PATTERNS: List[Tuple[str, str]] = [
    ("housing", r"\b(housing|accommodation|rental|temporary housing)\b"),
    ("household_goods", r"\b(household goods|shipment|moving|movers|freight)\b"),
    ("tuition", r"\b(tuition|education|schooling)\b"),
    ("home_leave", r"\b(home leave|home flight)\b"),
    ("mobility_premium", r"\b(mobility premium|expatriate premium)\b"),
    ("settling_in", r"\b(settling[- ]?in|relocation allowance)\b"),
    ("tax_equalization", r"\b(tax equalization|hypothetical tax)\b"),
    ("spouse_support", r"\b(spousal support|partner support)\b"),
]

CURRENCY_PATTERN = re.compile(r"\b(USD|EUR|GBP|CHF|CAD|AUD|JPY)\b", re.I)
UNIT_PATTERNS = [
    (r"\b(%|percent|percentage)\b", "%"),
    (r"\b(days?)\b", "days"),
    (r"\b(weeks?)\b", "weeks"),
    (r"\b(months?)\b", "months"),
    (r"\b(years?)\b", "years"),
    (r"\b(lbs?|kg)\b", "weight"),
]
NUMERIC_PATTERN = re.compile(
    r"(?:\b|\$|€|£)\s*([0-9]+(?:,[0-9]{3})*(?:\.[0-9]+)?)\s*(?:k|K|m|M|%|USD|EUR|GBP)?\b"
)
FREQUENCY_PATTERNS = [
    (r"\b(per year|annually|yearly)\b", "per_year"),
    (r"\b(per assignment|per move)\b", "per_assignment"),
    (r"\b(monthly|per month)\b", "monthly"),
    (r"\b(one[- ]?time|once)\b", "one_time"),
]
ASSIGNMENT_TYPE_TERMS = ["long_term", "short_term", "permanent", "commuter", "international"]
ASSIGNMENT_PATTERNS = [
    (r"\b(long[- ]?term|long term|lta)\b", "long_term"),
    (r"\b(short[- ]?term|short term|sta)\b", "short_term"),
    (r"\b(permanent transfer|permanent relocation)\b", "permanent"),
    (r"\b(commuter)\b", "commuter"),
]
FAMILY_TERMS = ["single", "married", "spouse", "dependents", "accompanying", "family"]
EXCLUSION_SIGNALS = ["exclusion", "excluded", "not covered", "ineligible", "does not apply", "out of scope"]
APPROVAL_SIGNALS = ["approval", "pre-approval", "prior approval", "requires approval", "hr approval"]
EVIDENCE_ITEMS = ["receipt", "receipts", "invoice", "invoices", "documentation", "proof", "quote", "estimates"]


def _extract_normalized_hints(
    text: str, is_table_row: bool, section_label: Optional[str]
) -> Dict[str, Any]:
    """
    Extract first-pass hints from clause text. All optional, non-authoritative.
    """
    lower = text.lower()
    hints: Dict[str, Any] = {}

    # candidate_benefit_key
    for key, pat in BENEFIT_KEY_PATTERNS:
        if re.search(pat, lower, re.I):
            hints["candidate_benefit_key"] = key
            break
    if not hints.get("candidate_benefit_key") and section_label:
        sl = section_label.lower()
        for key, _ in BENEFIT_KEY_PATTERNS:
            if key.replace("_", " ") in sl or key in sl:
                hints["candidate_benefit_key"] = key
                break

    # candidate_currency
    m = CURRENCY_PATTERN.search(text)
    if m:
        hints["candidate_currency"] = m.group(1).upper()

    # candidate_unit
    for pat, unit in UNIT_PATTERNS:
        if re.search(pat, lower):
            hints["candidate_unit"] = unit
            break

    # candidate_numeric_values
    nums = []
    for m in NUMERIC_PATTERN.finditer(text):
        s = m.group(1).replace(",", "")
        try:
            n = float(s)
            if 0 < n < 1e12:
                nums.append(n)
        except ValueError:
            pass
    if nums:
        hints["candidate_numeric_values"] = list(dict.fromkeys(nums))[:10]

    # candidate_frequency
    for pat, freq in FREQUENCY_PATTERNS:
        if re.search(pat, lower):
            hints["candidate_frequency"] = freq
            break

    # candidate_assignment_types
    ats = []
    for pat, at in ASSIGNMENT_PATTERNS:
        if re.search(pat, lower):
            ats.append(at)
    if ats:
        hints["candidate_assignment_types"] = list(dict.fromkeys(ats))

    # candidate_family_status_terms
    fts = [t for t in FAMILY_TERMS if t in lower]
    if fts:
        hints["candidate_family_status_terms"] = fts

    # candidate_exclusion_flag
    if any(s in lower for s in EXCLUSION_SIGNALS):
        hints["candidate_exclusion_flag"] = True

    # candidate_approval_flag
    if any(s in lower for s in APPROVAL_SIGNALS):
        hints["candidate_approval_flag"] = True

    # candidate_evidence_items
    evs = [e for e in EVIDENCE_ITEMS if e in lower]
    if evs:
        hints["candidate_evidence_items"] = list(dict.fromkeys(evs))

    return hints


def extract_lines_with_pages(
    data: bytes, mime_type: str
) -> Tuple[List[Dict[str, Any]], Optional[str]]:
    """
    Extract text with page numbers and table markers.
    Returns list of {"text": str, "page": int, "is_table_row": bool}.
    """
    if mime_type in ("application/pdf", "pdf") or (
        isinstance(mime_type, str) and "pdf" in mime_type.lower()
    ):
        return _extract_pdf_with_pages(data)
    if mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "docx",
    ) or (
        isinstance(mime_type, str)
        and ("word" in mime_type.lower() or "docx" in mime_type.lower())
    ):
        return _extract_docx_with_pages(data)
    return [], f"Unsupported mime type: {mime_type}"


def _extract_pdf_with_pages(data: bytes) -> Tuple[List[Dict[str, Any]], Optional[str]]:
    try:
        import pdfplumber
        import io
    except ImportError as exc:
        return [], f"pdfplumber required: {exc}"
    try:
        items: List[Dict[str, Any]] = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                for line in text.splitlines():
                    s = re.sub(r"\s+", " ", line.strip())
                    if s:
                        items.append({"text": s, "page": page_num, "is_table_row": False})
                for table in page.extract_tables() or []:
                    for row in table:
                        cells = [str(c).strip() if c else "" for c in row if c is not None]
                        if cells:
                            items.append({
                                "text": " | ".join(cells),
                                "page": page_num,
                                "is_table_row": True,
                            })
        return items, None
    except Exception as e:
        log.warning("pdf extraction with pages failed: %s", e, exc_info=True)
        return [], str(e)


def _extract_docx_with_pages(data: bytes) -> Tuple[List[Dict[str, Any]], Optional[str]]:
    try:
        from docx import Document
        import io
    except ImportError as exc:
        return [], f"python-docx required: {exc}"
    try:
        items: List[Dict[str, Any]] = []
        doc = Document(io.BytesIO(data))
        for p in doc.paragraphs:
            s = re.sub(r"\s+", " ", (p.text or "").strip())
            if s:
                items.append({"text": s, "page": 1, "is_table_row": False})
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text]
                if cells:
                    items.append({
                        "text": " | ".join(cells),
                        "page": 1,
                        "is_table_row": True,
                    })
        return items, None
    except Exception as e:
        log.warning("docx extraction with pages failed: %s", e, exc_info=True)
        return [], str(e)


def _is_heading(line: str) -> bool:
    """Detect heading patterns: numbering (1., 1.1, a), all caps, short lines."""
    s = line.strip()
    if not s or len(s) > 120:
        return False
    # Numbering: 1. 1.1 1.1.1 a) (1) I. II.
    if re.match(r"^(\d+\.)+\s", s) or re.match(r"^[a-zA-Z][\).]\s", s):
        return True
    if re.match(r"^\(\d+\)\s", s) or re.match(r"^[IVX]+\.\s", s, re.I):
        return True
    # All caps and short
    if s.isupper() and len(s) < 80:
        return True
    return False


def _is_section_name(line: str) -> Optional[str]:
    """Check if line matches a known section heading."""
    lower = line.lower().strip()
    for heading in SECTION_HEADINGS:
        if heading in lower or lower == heading:
            return heading
    return None


def _classify_clause(text: str, is_table_row: bool) -> Tuple[str, float]:
    """
    Rule-based clause type classification. Returns (clause_type, confidence).
    """
    lower = text.lower()
    best = (CLAUSE_UNKNOWN, 0.3)

    for signals, ctype in CLAUSE_TYPE_SIGNALS:
        for sig in signals:
            if sig in lower:
                # Benefit is default for table rows with amounts/percentages
                if ctype == CLAUSE_BENEFIT or (is_table_row and " | " in text):
                    conf = 0.75
                else:
                    conf = 0.8
                if conf > best[1]:
                    best = (ctype, conf)

    # Table rows often encode benefits
    if is_table_row and " | " in text and best[0] == CLAUSE_UNKNOWN:
        if any(x in lower for x in ["usd", "eur", "%", "limit", "allowance", "cover"]):
            best = (CLAUSE_BENEFIT, 0.7)

    return best


def segment_into_clauses(
    items: List[Dict[str, Any]], raw_text: Optional[str] = None
) -> List[Dict[str, Any]]:
    """
    Segment extracted lines into clause blocks.
    items: list of {text, page, is_table_row} from extract_lines_with_pages.
    Returns list of clause dicts ready for DB insert.
    """
    if not items:
        return []

    clauses: List[Dict[str, Any]] = []
    current_section = None
    current_path: List[str] = []
    buffer: List[Dict[str, Any]] = []
    current_page_start = items[0].get("page", 1)

    def flush_buffer(title: Optional[str] = None, clause_type: Optional[str] = None):
        if not buffer:
            return
        texts = [b["text"] for b in buffer]
        pages = [b["page"] for b in buffer]
        is_table = any(b.get("is_table_row") for b in buffer)
        raw = "\n".join(texts)
        ctype, conf = _classify_clause(raw, is_table)
        if clause_type:
            ctype = clause_type
            conf = 0.9
        section_path = " > ".join(current_path) if current_path else None
        hints = _extract_normalized_hints(raw, is_table, current_section)
        clauses.append({
            "section_label": current_section,
            "section_path": section_path,
            "clause_type": ctype,
            "title": title,
            "raw_text": raw,
            "normalized_hint_json": hints if hints else None,
            "source_page_start": min(pages) if pages else None,
            "source_page_end": max(pages) if pages else None,
            "source_anchor": texts[0][:100] if texts else None,
            "confidence": conf,
        })
        buffer.clear()

    i = 0
    while i < len(items):
        item = items[i]
        text = item.get("text", "").strip()
        page = item.get("page", 1)
        is_table = item.get("is_table_row", False)

        # Section heading
        section = _is_section_name(text)
        if section:
            flush_buffer()
            current_section = section
            # Use first heading word as path component
            path_part = section.split()[0] if section else None
            if path_part and (not current_path or current_path[-1] != path_part):
                current_path.append(path_part)
            buffer.append(item)
            flush_buffer(title=text, clause_type=None)
            current_page_start = page
            i += 1
            continue

        # Numbered heading (1.1, 2.3.1, etc.)
        if _is_heading(text) and not is_table:
            flush_buffer()
            match = re.match(r"^((?:\d+\.)+\d*)\s*(.+)", text)
            if match:
                num, rest = match.groups()
                current_path = num.rstrip(".").split(".")
                current_section = rest.strip()
                buffer.append(item)
                flush_buffer(title=text)
            else:
                buffer.append(item)
            i += 1
            continue

        # Table row: emit as separate clause for traceability
        if is_table:
            flush_buffer()
            ctype, conf = _classify_clause(text, True)
            section_path = " > ".join(current_path) if current_path else None
            hints = _extract_normalized_hints(text, True, current_section)
            clauses.append({
                "section_label": current_section,
                "section_path": section_path,
                "clause_type": ctype,
                "title": None,
                "raw_text": text,
                "normalized_hint_json": hints if hints else None,
                "source_page_start": page,
                "source_page_end": page,
                "source_anchor": text[:100],
                "confidence": conf,
            })
            i += 1
            continue

        # Bullet or continuation
        buffer.append(item)
        i += 1

    flush_buffer()
    return clauses


def segment_document_from_raw_text(
    raw_text: str, mime_type: str, data: Optional[bytes] = None
) -> Tuple[List[Dict[str, Any]], Optional[str]]:
    """
    Segment a document. Prefer data for page-aware extraction; fallback to raw_text lines.
    Returns (clauses, error).
    """
    if data:
        items, err = extract_lines_with_pages(data, mime_type)
        if err:
            return [], err
    else:
        # Fallback: treat raw_text as single page
        lines = [l.strip() for l in raw_text.splitlines() if l.strip()]
        items = [{"text": l, "page": 1, "is_table_row": " | " in l and len(l) > 20}
                 for l in lines]
    return segment_into_clauses(items), None
