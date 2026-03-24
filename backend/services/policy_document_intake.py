"""
Policy document intake pipeline (INGEST stage): parse PDF/DOCX, classify, extract metadata.

Pipeline layers (see docs/policy/metadata-vs-decision-layer.md):
- **Layer 1 — document metadata / structure:** Everything in this module (classification,
  extracted_metadata, heuristic flags, mentioned_* terms) is descriptive only. It must not be
  treated as binding coverage for employees.
- **Layer 2 — decision output:** Produced later by normalization + publish + resolution
  (policy_benefit_rules, resolved_assignment_policy_benefits, etc.).

Pipeline stage (3-stage model):
- Ingest (this module): store file in blob storage, extract raw text, classify document type/scope,
  extract metadata (title, version, effective date), and optionally segment into clauses.
  Called from the upload endpoint after the file is stored. Output: policy_documents row + clauses.
- Reprocess: re-run extraction and clause segmentation from the stored file without re-uploading.
  Used when extraction logic or segmentation improves, or to fix failed extraction.
- Normalize: transform clauses into structured policy objects (company_policies, policy_versions,
  benefit_rules, exclusions, etc.). Separate so HR can reprocess without overwriting normalized edits.
"""
from __future__ import annotations

import hashlib
import io
import json
import logging
import re
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

log = logging.getLogger(__name__)

# Processing statuses
STATUS_UPLOADED = "uploaded"
STATUS_TEXT_EXTRACTED = "text_extracted"
STATUS_CLASSIFIED = "classified"
STATUS_NORMALIZED = "normalized"
STATUS_REVIEW_REQUIRED = "review_required"
STATUS_APPROVED = "approved"
STATUS_FAILED = "failed"

# Document types
DOC_TYPE_ASSIGNMENT_POLICY = "assignment_policy"
DOC_TYPE_POLICY_SUMMARY = "policy_summary"
DOC_TYPE_SUMMARY_TABLE = "summary_table"
DOC_TYPE_COMPACT_BENEFIT_MATRIX = "compact_benefit_matrix"
DOC_TYPE_TAX_POLICY = "tax_policy"
DOC_TYPE_COUNTRY_ADDENDUM = "country_addendum"
DOC_TYPE_UNKNOWN = "unknown"

# Policy scopes
SCOPE_GLOBAL = "global"
SCOPE_LONG_TERM = "long_term_assignment"
SCOPE_SHORT_TERM = "short_term_assignment"
SCOPE_TAX_EQUALIZATION = "tax_equalization"
SCOPE_MIXED = "mixed"
SCOPE_UNKNOWN = "unknown"

BENEFIT_CATEGORY_KEYWORDS: Dict[str, List[str]] = {
    "housing": ["housing", "temporary housing", "rental", "accommodation", "deposit"],
    "movers": ["shipment", "household goods", "movers", "moving", "freight"],
    "schools": ["education", "school", "tuition", "childcare"],
    "immigration": ["visa", "immigration", "work permit", "residence permit"],
    "travel": ["travel", "flight", "airfare", "scouting trip", "home leave"],
    "settling_in": ["settling-in", "settling in", "allowance"],
    "tax": ["tax assistance", "tax equalization", "hypothetical tax", "tax"],
    "spouse": ["spousal support", "spouse", "partner support"],
    "integration": ["language", "cultural", "integration", "training"],
    "repatriation": ["repatriation", "return shipment", "return travel"],
}

ASSIGNMENT_TYPE_PATTERNS = [
    ("long_term", ["long-term assignment", "long term assignment", "lta", "lt assignment", "long-term", "long term"]),
    ("short_term", ["short-term assignment", "short term assignment", "sta", "st assignment", "short-term", "short term"]),
    ("permanent", ["permanent transfer", "permanent relocation", "permanent move"]),
    ("commuter", ["commuter", "commuter assignment"]),
    ("extended_business_trip", ["extended business trip", "ebt"]),
    ("international", ["international assignment", "global assignment"]),
]

FAMILY_STATUS_TERMS = [
    "single", "married", "unmarried", "spouse", "accompanying spouse", "trailing spouse",
    "dependents", "dependent children", "family", "accompanying family",
    "domestic partner", "partner", "household",
]

UNIT_PATTERNS = [
    r"\b(usd|eur|gbp|chf|cad|aud|jpy)\b",
    r"\b(%|percent|percentage)\b",
    r"\b(days?|weeks?|months?|years?)\b",
    r"\b(lbs?|kg)\b",
]

EXCLUSION_SIGNALS = [
    "exclusion", "excluded", "excluding", "not covered", "ineligible",
    "does not apply", "out of scope", "outside policy", "non-covered",
]

APPROVAL_SIGNALS = [
    "approval", "pre-approval", "pre approval", "hr approval", "manager approval",
    "requires approval", "prior approval", "approved by", "sign-off",
]

EVIDENCE_SIGNALS = [
    "evidence", "receipt", "invoice", "documentation required", "proof of",
    "supporting documentation", "submit receipts", "receipts required",
]


def compute_checksum(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def extract_text_from_bytes(data: bytes, mime_type: str) -> Tuple[List[str], Optional[str]]:
    """
    Extract text from PDF or DOCX. Returns (lines, error).
    """
    if mime_type in ("application/pdf", "pdf") or (isinstance(mime_type, str) and "pdf" in mime_type.lower()):
        return _extract_text_from_pdf(data)
    if mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "docx",
    ) or (isinstance(mime_type, str) and "word" in mime_type.lower() or "docx" in mime_type.lower()):
        return _extract_text_from_docx(data)
    return [], f"Unsupported mime type: {mime_type}"


def _normalize_lines(lines: List[str]) -> List[str]:
    cleaned = []
    for line in lines:
        if not line:
            continue
        s = re.sub(r"\s+", " ", line.strip())
        if s:
            cleaned.append(s)
    return cleaned


def _extract_text_from_docx(data: bytes) -> Tuple[List[str], Optional[str]]:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        return [], f"python-docx required: {exc}"
    try:
        doc = Document(io.BytesIO(data))
        lines: List[str] = []
        for p in doc.paragraphs:
            lines.append(p.text or "")
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text]
                if cells:
                    lines.append(" | ".join(cells))
        return _normalize_lines(lines), None
    except Exception as e:
        log.warning("docx extraction failed: %s", e, exc_info=True)
        return [], str(e)


def _extract_text_from_pdf(data: bytes) -> Tuple[List[str], Optional[str]]:
    try:
        import pdfplumber  # type: ignore
    except ImportError as exc:
        return [], f"pdfplumber required: {exc}"
    try:
        lines: List[str] = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                lines.extend(text.splitlines())
                for table in page.extract_tables() or []:
                    for row in table:
                        cells = [str(c) if c else "" for c in row if c is not None]
                        if cells:
                            lines.append(" | ".join(cells))
        return _normalize_lines(lines), None
    except Exception as e:
        log.warning("pdf extraction failed: %s", e, exc_info=True)
        return [], str(e)


def classify_document(lines: List[str], request_id: Optional[str] = None) -> Tuple[str, str, bool]:
    """
    Rule-based classifier. Returns (document_type, policy_scope, needs_review).
    """
    text = "\n".join(lines).lower()
    doc_type = DOC_TYPE_UNKNOWN
    scope = SCOPE_UNKNOWN
    needs_review = True

    # Tax policy signals
    tax_signals = [
        "tax equalization",
        "hypothetical tax",
        "hypothetical social security",
        "hypothetical tax calculation",
        "tax protection",
    ]
    if any(s in text for s in tax_signals):
        doc_type = DOC_TYPE_TAX_POLICY
        scope = SCOPE_TAX_EQUALIZATION
        needs_review = False
        log.info(
            "request_id=%s classify_document: tax_policy (tax signals)",
            request_id or "",
        )
        return (doc_type, scope, needs_review)

    # Policy summary signals
    if "long term assignment policy summary" in text or "lta policy summary" in text:
        doc_type = DOC_TYPE_POLICY_SUMMARY
        scope = SCOPE_LONG_TERM
        needs_review = False
        log.info(
            "request_id=%s classify_document: policy_summary (LTA summary)",
            request_id or "",
        )
        return (doc_type, scope, needs_review)

    # Assignment policy: International Assignment Management + mobility sections
    assignment_signals = [
        "international assignment management",
        "international assignment policy",
        "global mobility policy",
    ]
    mobility_sections = [
        "mobility premium",
        "home leave",
        "moving services",
        "household goods",
        "relocation allowance",
        "housing allowance",
    ]
    has_assignment_title = any(s in text for s in assignment_signals)
    has_mobility_content = sum(1 for s in mobility_sections if s in text) >= 2
    if has_assignment_title and has_mobility_content:
        doc_type = DOC_TYPE_ASSIGNMENT_POLICY
        scope = SCOPE_GLOBAL
        needs_review = False
        log.info(
            "request_id=%s classify_document: assignment_policy (title + mobility sections)",
            request_id or "",
        )
        return (doc_type, scope, needs_review)

    # Fallback: if "international" + "assignment" present
    if "international" in text and "assignment" in text:
        doc_type = DOC_TYPE_ASSIGNMENT_POLICY
        scope = SCOPE_MIXED
        needs_review = True
        log.info("request_id=%s classify_document: assignment_policy (fallback, needs_review)", request_id or "")
        return (doc_type, scope, needs_review)

    # Country addendum: country name + addendum
    country_addendum = re.search(r"(addendum|annex|appendix)\s+.*\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b", text, re.I)
    if country_addendum and any(c in text for c in ["country", "local", "host"]):
        doc_type = DOC_TYPE_COUNTRY_ADDENDUM
        scope = SCOPE_UNKNOWN
        needs_review = True
        log.info("request_id=%s classify_document: country_addendum", request_id or "")
        return (doc_type, scope, needs_review)

    log.info("request_id=%s classify_document: unknown (no rules matched)", request_id or "")
    return (doc_type, scope, needs_review)


def _empty_metadata() -> Dict[str, Any]:
    """Canonical empty extracted_metadata schema."""
    return {
        "detected_title": None,
        "detected_version": None,
        "detected_effective_date": None,
        "mentioned_assignment_types": [],
        "mentioned_family_status_terms": [],
        "mentioned_benefit_categories": [],
        "mentioned_units": [],
        "likely_table_heavy": False,
        "likely_country_addendum": False,
        "likely_tax_specific": False,
        "likely_contains_exclusions": False,
        "likely_contains_approval_rules": False,
        "likely_contains_evidence_rules": False,
    }


def normalize_extracted_metadata(raw: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Normalize extracted_metadata to the current schema.
    Backward-compatible with legacy fields (policy_title, version, contains_tables, etc.).
    """
    base = _empty_metadata()
    if not raw or not isinstance(raw, dict):
        return base

    # Map legacy keys to new schema
    base["detected_title"] = raw.get("detected_title") or raw.get("policy_title")
    base["detected_version"] = raw.get("detected_version") or raw.get("version")
    base["detected_effective_date"] = raw.get("detected_effective_date") or raw.get("effective_date")
    base["mentioned_assignment_types"] = raw.get("mentioned_assignment_types")
    if base["mentioned_assignment_types"] is None:
        leg = raw.get("assignment_types_mentioned")
        base["mentioned_assignment_types"] = leg if isinstance(leg, list) else []

    base["mentioned_family_status_terms"] = raw.get("mentioned_family_status_terms")
    if base["mentioned_family_status_terms"] is None:
        base["mentioned_family_status_terms"] = []

    base["mentioned_benefit_categories"] = raw.get("mentioned_benefit_categories")
    if base["mentioned_benefit_categories"] is None:
        leg = raw.get("benefit_categories_mentioned")
        base["mentioned_benefit_categories"] = leg if isinstance(leg, list) else []

    base["mentioned_units"] = raw.get("mentioned_units")
    if base["mentioned_units"] is None:
        base["mentioned_units"] = []

    base["likely_table_heavy"] = bool(
        raw.get("likely_table_heavy")
        if "likely_table_heavy" in raw
        else raw.get("contains_tables", False)
    )
    base["likely_country_addendum"] = bool(raw.get("likely_country_addendum", False))
    base["likely_tax_specific"] = bool(raw.get("likely_tax_specific", False))
    base["likely_contains_exclusions"] = bool(raw.get("likely_contains_exclusions", False))
    base["likely_contains_approval_rules"] = bool(raw.get("likely_contains_approval_rules", False))
    base["likely_contains_evidence_rules"] = bool(raw.get("likely_contains_evidence_rules", False))

    return base


def extract_metadata(lines: List[str]) -> Dict[str, Any]:
    """
    Rule-based, deterministic extraction of structured metadata.
    Returns canonical schema for extracted_metadata.
    """
    meta = _empty_metadata()
    text_lower = "\n".join(lines).lower()
    text_full = "\n".join(lines)

    # detected_title: first substantial line containing "policy"
    for line in lines[:40]:
        s = line.strip()
        if s and "policy" in s.lower() and 5 < len(s) < 150:
            meta["detected_title"] = s
            break

    # detected_version
    for line in lines[:60]:
        m = re.search(r"(?:version|v\.?)\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)", line, re.I)
        if m:
            meta["detected_version"] = m.group(1)
            break

    # detected_effective_date
    for line in lines[:80]:
        m = re.search(r"(?:effective|valid from)\s*[:\-]?\s*([0-9]{4}-[0-9]{2}-[0-9]{2})", line, re.I)
        if m:
            meta["detected_effective_date"] = m.group(1)
            break
        if not meta["detected_effective_date"]:
            m = re.search(r"([0-9]{1,2}[/\-][0-9]{1,2}[/\-][0-9]{4})", line)
            if m:
                d = m.group(1)
                parts = re.split(r"[/\-]", d)
                if len(parts) == 3:
                    y, mo, day = (
                        (parts[2], parts[0], parts[1])
                        if len(parts[2]) == 4
                        else (parts[2], parts[1], parts[0])
                    )
                    meta["detected_effective_date"] = f"{y}-{mo.zfill(2)}-{day.zfill(2)}"
                    break
        if not meta["detected_effective_date"]:
            m = re.search(r"([0-9]{4})-([0-9]{2})-([0-9]{2})", line)
            if m:
                meta["detected_effective_date"] = m.group(0)
                break

    # mentioned_assignment_types
    for key, patterns in ASSIGNMENT_TYPE_PATTERNS:
        if any(p in text_lower for p in patterns) and key not in meta["mentioned_assignment_types"]:
            meta["mentioned_assignment_types"].append(key)

    # mentioned_family_status_terms
    for term in FAMILY_STATUS_TERMS:
        if term in text_lower and term not in meta["mentioned_family_status_terms"]:
            meta["mentioned_family_status_terms"].append(term)

    # mentioned_benefit_categories
    for cat, keywords in BENEFIT_CATEGORY_KEYWORDS.items():
        if any(k in text_lower for k in keywords) and cat not in meta["mentioned_benefit_categories"]:
            meta["mentioned_benefit_categories"].append(cat)

    # mentioned_units (currencies, %, time units)
    seen: set = set()
    for pat in UNIT_PATTERNS:
        for m in re.finditer(pat, text_lower, re.I):
            u = m.group(1).lower()
            if u not in seen:
                seen.add(u)
                meta["mentioned_units"].append(u)

    # likely_table_heavy: many pipe-separated lines
    table_like = sum(1 for ln in lines if " | " in ln and len(ln) > 20)
    meta["likely_table_heavy"] = table_like >= 3

    # likely_country_addendum
    addendum_match = re.search(
        r"(addendum|annex|appendix)\s+.*\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b",
        text_full,
        re.I,
    )
    meta["likely_country_addendum"] = bool(
        addendum_match and any(c in text_lower for c in ["country", "local", "host"])
    )

    # likely_tax_specific
    meta["likely_tax_specific"] = any(
        s in text_lower
        for s in [
            "tax equalization",
            "hypothetical tax",
            "hypothetical social security",
            "tax protection",
        ]
    )

    # likely_contains_exclusions
    meta["likely_contains_exclusions"] = any(s in text_lower for s in EXCLUSION_SIGNALS)

    # likely_contains_approval_rules
    meta["likely_contains_approval_rules"] = any(s in text_lower for s in APPROVAL_SIGNALS)

    # likely_contains_evidence_rules
    meta["likely_contains_evidence_rules"] = any(s in text_lower for s in EVIDENCE_SIGNALS)

    return meta


def process_uploaded_document(
    data: bytes,
    mime_type: str,
    filename: str,
    request_id: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Full intake pipeline: extract text, classify, extract metadata.
    Returns dict with raw_text, detected_document_type, detected_policy_scope,
    extracted_metadata, processing_status, extraction_error.
    """
    result: Dict[str, Any] = {
        "raw_text": None,
        "detected_document_type": DOC_TYPE_UNKNOWN,
        "detected_policy_scope": SCOPE_UNKNOWN,
        "extracted_metadata": {},
        "processing_status": STATUS_UPLOADED,
        "extraction_error": None,
        "version_label": None,
        "effective_date": None,
    }

    try:
        lines, err = extract_text_from_bytes(data, mime_type)
        if err:
            result["processing_status"] = STATUS_FAILED
            result["extraction_error"] = err
            log.warning("request_id=%s process_uploaded_document text_extract failed: %s", request_id, err)
            return result

        raw_text = "\n".join(lines) if lines else ""
        result["raw_text"] = raw_text
        result["processing_status"] = STATUS_TEXT_EXTRACTED

        if not lines:
            result["processing_status"] = STATUS_FAILED
            result["extraction_error"] = "No text extracted from document"
            return result

        doc_type, scope, needs_review = classify_document(lines, request_id=request_id)
        result["detected_document_type"] = doc_type
        result["detected_policy_scope"] = scope
        result["processing_status"] = STATUS_CLASSIFIED if not needs_review else STATUS_REVIEW_REQUIRED

        meta = extract_metadata(lines)
        result["extracted_metadata"] = meta
        result["version_label"] = meta.get("detected_version")
        result["effective_date"] = meta.get("detected_effective_date")

    except Exception as e:
        result["processing_status"] = STATUS_FAILED
        result["extraction_error"] = str(e)
        log.warning(
            "request_id=%s process_uploaded_document failed: %s",
            request_id,
            e,
            exc_info=True,
        )

    return result
